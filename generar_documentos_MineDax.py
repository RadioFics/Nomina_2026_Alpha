"""
generar_documentos_MineDax.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Genera dos archivos .docx para el proyecto MineDax 2026:
  1. InformeTecnico_Requerimientos_MineDax2026.docx
  2. MensajeInformativo_Requerimientos_MineDax2026.docx

EJECUTAR DESDE PowerShell o CMD (en la misma carpeta):
  python generar_documentos_MineDax.py

Si falta python-docx, el script lo instala automáticamente.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import sys
import subprocess
import os

# Auto-install python-docx
try:
    from docx import Document
except ImportError:
    print("Instalando python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Output same folder as this script
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── Colores ──────────────────────────────────────────────────────────────────
ACCENT_BLUE = "1F4E79"
LIGHT_BLUE  = "2E75B6"
MID_BLUE    = "D6E4F0"
LIGHT_GRAY  = "F2F2F2"
DARK_TEXT   = "1A1A1A"
WHITE       = "FFFFFF"
GREEN       = "375623"
LT_GREEN    = "E2EFDA"
ORANGE      = "833C00"
LT_ORANGE   = "FCE4D6"
LT_GRAY2    = "F5F5F5"


def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    for e in tcPr.findall(qn('w:shd')):
        tcPr.remove(e)
    tcPr.append(shd)


def set_cell_borders(cell, color="AAAAAA", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    for e in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(e)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    for e in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(e)
    tcPr.append(tcMar)


def set_cell_valign(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    for e in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(e)
    tcPr.append(vAlign)


def set_col_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    for e in tcPr.findall(qn('w:tcW')):
        tcPr.remove(e)
    tcPr.append(tcW)


def set_table_width(table, width_dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    for e in tblPr.findall(qn('w:tblW')):
        tblPr.remove(e)
    tblPr.append(tblW)


def add_run(para, text, bold=False, italic=False, color=DARK_TEXT, size_pt=11, font="Arial"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    r, g, b = hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def add_paragraph(doc, text, bold=False, italic=False, color=DARK_TEXT, size_pt=11,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        add_run(p, text, bold=bold, italic=italic, color=color, size_pt=size_pt)
    return p


def add_divider(doc, color=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def add_section_label(doc, text, fill=MID_BLUE, border_color=LIGHT_BLUE, text_color=ACCENT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '20')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    add_run(p, "  " + text, bold=True, color=text_color, size_pt=12)
    return p


def add_bullet(doc, text, color=DARK_TEXT):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, color=color, size_pt=10)
    return p


def add_heading1(doc, text, color=ACCENT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '2')
    bot.set(qn('w:color'), LIGHT_BLUE)
    pBdr.append(bot)
    pPr.append(pBdr)
    add_run(p, text, bold=True, color=color, size_pt=16)
    return p


def add_heading2(doc, text, color=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, text, bold=True, color=color, size_pt=13)
    return p


def insert_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p._p.append(r)
    return p


def make_header_row(table, labels, widths, fill=LIGHT_BLUE, text_color=WHITE):
    row = table.add_row()
    for i, lbl in enumerate(labels):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, fill)
        set_cell_borders(cell, color=LIGHT_BLUE, size=4)
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        set_cell_valign(cell)
        set_col_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, lbl, bold=True, color=text_color, size_pt=10)
    return row


def make_data_row(table, cells_text, widths, shade=WHITE):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, shade)
        set_cell_borders(cell, color="AAAAAA", size=4)
        set_cell_margins(cell)
        set_cell_valign(cell)
        set_col_width(cell, widths[i])
        p = cell.paragraphs[0]
        add_run(p, text, color=DARK_TEXT, size_pt=9)
    return row


def add_callout(doc, label, text, fill=MID_BLUE, border_color=LIGHT_BLUE, text_color=ACCENT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 9026)
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_bg(cell, fill)
    set_col_width(cell, 9026)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, sz, color in [('top', 6, border_color), ('bottom', 6, border_color),
                              ('left', 20, border_color), ('right', 6, border_color)]:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    for e in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(e)
    tcPr.append(tcBorders)
    p_label = cell.paragraphs[0]
    p_label.paragraph_format.space_after = Pt(3)
    add_run(p_label, label, bold=True, color=text_color, size_pt=10.5)
    p_text = cell.add_paragraph()
    p_text.paragraph_format.space_after = Pt(0)
    add_run(p_text, text, color=DARK_TEXT, size_pt=10)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)
    return table


def set_page_margins(doc):
    section = doc.sections[0]
    section.top_margin    = Twips(1440)
    section.bottom_margin = Twips(1440)
    section.left_margin   = Twips(1300)
    section.right_margin  = Twips(1300)
    section.page_width    = Twips(11906)
    section.page_height   = Twips(16838)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 1: INFORME TÉCNICO
# ═══════════════════════════════════════════════════════════════════════════════
def build_informe_tecnico():
    doc = Document()
    set_page_margins(doc)

    # PORTADA
    p_b = doc.add_paragraph()
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_b.paragraph_format.space_before = Pt(30)
    p_b.paragraph_format.space_after = Pt(0)
    pPr = p_b._p.get_or_add_pPr()
    s = OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'), ACCENT_BLUE)
    pPr.append(s)
    add_run(p_b, "MineDax 2026", bold=True, color=WHITE, size_pt=30)

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s.paragraph_format.space_before = Pt(0)
    p_s.paragraph_format.space_after = Pt(10)
    pPr2 = p_s._p.get_or_add_pPr()
    s2 = OxmlElement('w:shd'); s2.set(qn('w:val'),'clear'); s2.set(qn('w:color'),'auto'); s2.set(qn('w:fill'), ACCENT_BLUE)
    pPr2.append(s2)
    add_run(p_s, "Plataforma de Gestión en Línea", color="BDD7EE", size_pt=16)

    doc.add_paragraph()
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_t, "INFORME TÉCNICO DE REQUERIMIENTOS FUNCIONALES", bold=True, color=ACCENT_BLUE, size_pt=18)

    p_s2 = doc.add_paragraph()
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s2.paragraph_format.space_after = Pt(4)
    add_run(p_s2, "Base de Datos  ·  Interfaz Web  ·  Formularios  ·  Gestor de Interfaz", italic=True, color="555555", size_pt=12)

    add_divider(doc)

    meta = doc.add_table(rows=0, cols=2)
    set_table_width(meta, 9026)
    for label, value, shade in [
        ("Versión", "1.0 - Alfa", LIGHT_GRAY),
        ("Fecha de elaboración", "Mayo 2026", WHITE),
        ("Clasificación", "Confidencial – Uso interno", LIGHT_GRAY),
        ("Proyecto", "MineDax 2026", WHITE),
        ("Línea de despliegue", "Línea Privada en Línea", LIGHT_GRAY),
        ("Dirigido a", "Equipo de Desarrollo / Mesa de Ayuda", WHITE),
    ]:
        make_data_row(meta, [label, value], [2500, 6526], shade=shade)

    insert_page_break(doc)

    # 1. INTRODUCCIÓN
    add_heading1(doc, "1. Introducción")
    add_paragraph(doc, "El presente informe técnico tiene por objeto especificar de manera formal y exhaustiva los requerimientos funcionales necesarios para el correcto despliegue y operación de MineDax 2026, una plataforma de gestión en línea diseñada para operar bajo una línea privada de acceso. El documento sirve como guía de referencia para el equipo de desarrollo y el área de mesa de ayuda, con el fin de garantizar que cada componente del sistema sea implementado, configurado y mantenido según los estándares técnicos definidos.")
    add_paragraph(doc, "MineDax 2026 está concebido como un sistema modular que integra tres capas funcionales principales: la capa de datos (base de datos en línea y persistente), la capa de presentación (interfaz web y gestor de interfaz) y la capa de captura de información (formularios independientes y remotos). La correcta interoperabilidad entre estas capas es fundamental para garantizar la continuidad del servicio, la integridad de los datos y la disponibilidad de la plataforma para los usuarios autorizados.")
    add_divider(doc)

    # 2. DESCRIPCIÓN GENERAL
    add_heading1(doc, "2. Descripción General del Sistema")
    add_heading2(doc, "2.1 Propósito del Sistema")
    add_paragraph(doc, "MineDax 2026 es una plataforma web de gestión centralizada orientada a la administración de información operativa, documental y de usuarios dentro de una organización. Su objetivo principal es consolidar en un único entorno digital los procesos de registro, consulta, generación y almacenamiento de datos, eliminando la dependencia de procesos manuales o sistemas desconectados.")
    add_heading2(doc, "2.2 Alcance del Despliegue")
    add_paragraph(doc, "El sistema operará a través de una línea privada con acceso controlado, garantizando que únicamente los usuarios autorizados puedan interactuar con la plataforma. Esto implica la necesidad de un entorno de servidor dedicado o cloud privado, con mecanismos de autenticación robustos y gestión de roles diferenciada.")
    add_heading2(doc, "2.3 Lenguaje y Tecnologías Base")
    add_paragraph(doc, "La plataforma está construida con el siguiente stack tecnológico:")
    t = doc.add_table(rows=0, cols=3)
    set_table_width(t, 9026)
    make_header_row(t, ["Capa", "Tecnología / Lenguaje", "Justificación"], [2800, 3200, 3026])
    for i, row in enumerate([
        ("Frontend (Interfaz Web)", "React.js / Next.js", "Rendimiento, componentes reutilizables, SSR"),
        ("Backend (API/Lógica)", "Node.js + Express o Django (Python)", "API REST, escalabilidad, amplia comunidad"),
        ("Base de datos", "PostgreSQL / MySQL", "Relacional, ACID, soporte para consultas complejas"),
        ("Formularios remotos", "HTML5 + JS independiente / React Forms", "Autonomía del cliente sin dependencia del servidor principal"),
        ("Autenticación", "JWT + OAuth 2.0 / LDAP", "Seguridad estándar, soporte multiusuario"),
        ("Alojamiento / Infraestructura", "Servidor VPS privado / Cloud privado", "Control, privacidad, SLA garantizado"),
    ]):
        make_data_row(t, list(row), [2800, 3200, 3026], shade=WHITE if i%2==0 else LIGHT_GRAY)
    add_divider(doc)

    # 3. BASE DE DATOS
    insert_page_break(doc)
    add_heading1(doc, "3. Requerimientos de la Base de Datos")
    add_paragraph(doc, "La base de datos constituye el núcleo de persistencia y confiabilidad del sistema. Todos los datos generados por la plataforma —formularios, registros de usuario, transacciones, auditorías— deben quedar almacenados de forma permanente, consistente e íntegra, garantizando disponibilidad continua.")
    add_heading2(doc, "3.1 Requerimientos Funcionales de la Base de Datos")

    add_section_label(doc, "RF-DB-01 · Almacenamiento persistente y continuo")
    add_paragraph(doc, "La base de datos debe mantenerse activa y disponible de manera continua (24/7), con capacidad para recibir operaciones de escritura y lectura en tiempo real. Se requiere que todas las transacciones cumplan con las propiedades ACID.")
    for txt in ["Motor recomendado: PostgreSQL 15+ o MySQL 8+",
                "Modo de operación: persistente en servidor dedicado o instancia cloud privada",
                "Backups automáticos: diarios completos + incrementales por hora",
                "Replicación: al menos una réplica de lectura para alta disponibilidad"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-DB-02 · Estructura de tablas y entidades principales")
    add_paragraph(doc, "El esquema de base de datos deberá contemplar, como mínimo, las siguientes entidades:")
    t2 = doc.add_table(rows=0, cols=3)
    set_table_width(t2, 9026)
    make_header_row(t2, ["Entidad / Tabla", "Descripción", "Campos Clave"], [2600, 3000, 3426])
    for i, row in enumerate([
        ("usuarios", "Registro de todos los usuarios del sistema", "id, nombre, correo, rol, estado, fecha_creacion"),
        ("sesiones", "Control de sesiones activas y expiradas", "id, usuario_id, token, ip, fecha_inicio, fecha_fin"),
        ("formularios", "Metadata de formularios creados/recibidos", "id, tipo, nombre, fecha, estado, creado_por"),
        ("respuestas_formulario", "Datos capturados por cada formulario", "id, formulario_id, campo, valor, timestamp"),
        ("registros_operacion", "Historial de operaciones del sistema", "id, usuario_id, accion, modulo, timestamp, ip"),
        ("configuracion_sistema", "Parámetros de configuración globales", "clave, valor, descripcion, ultima_modificacion"),
        ("documentos", "Archivos adjuntos y documentos gestionados", "id, nombre, ruta, tipo, tamano, subido_por, fecha"),
    ]):
        make_data_row(t2, list(row), [2600, 3000, 3426], shade=WHITE if i%2==0 else LIGHT_GRAY)

    add_section_label(doc, "RF-DB-03 · Gestión de usuarios y roles")
    add_paragraph(doc, "El sistema de usuarios en base de datos debe soportar múltiples niveles de acceso:")
    t3 = doc.add_table(rows=0, cols=3)
    set_table_width(t3, 9026)
    make_header_row(t3, ["Tipo de Usuario", "Rol", "Permisos en BD"], [2000, 2800, 4226])
    for i, row in enumerate([
        ("Superadministrador", "SUPERADMIN", "CRUD completo sobre todas las tablas, gestión de esquema"),
        ("Administrador", "ADMIN", "CRUD sobre tablas de negocio, sin acceso a esquema"),
        ("Operador", "OPERATOR", "INSERT y SELECT en tablas asignadas"),
        ("Consultor", "VIEWER", "SELECT únicamente sobre vistas autorizadas"),
        ("Formulario remoto", "FORM_USER", "INSERT en tablas de respuestas, sin acceso a interfaz"),
    ]):
        make_data_row(t3, list(row), [2000, 2800, 4226], shade=WHITE if i%2==0 else LIGHT_GRAY)

    add_section_label(doc, "RF-DB-04 · Sincronización y conectividad remota")
    add_paragraph(doc, "La base de datos debe estar expuesta a través de una capa API segura que permita la recepción de datos desde formularios remotos.")
    for txt in ["Endpoint de API dedicado para recepción de formularios remotos",
                "Validación de datos en capa API antes de persistir en BD",
                "Cola de mensajes (ej. RabbitMQ o similar) para garantizar entrega en caso de intermitencia",
                "Registro de auditoría de todas las operaciones de escritura"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-DB-05 · Seguridad y cifrado")
    for txt in ["Conexiones cifradas mediante TLS/SSL obligatorio",
                "Contraseñas de usuarios almacenadas con hash bcrypt (factor >= 12)",
                "Datos sensibles cifrados en reposo con AES-256",
                "Acceso a BD restringido por IP y credenciales de servicio, no expuesta a internet directamente",
                "Política de retención y eliminación segura de datos"]:
        add_bullet(doc, txt)
    add_divider(doc)

    # 4. INTERFAZ WEB
    insert_page_break(doc)
    add_heading1(doc, "4. Requerimientos de la Interfaz Web")
    add_paragraph(doc, "La interfaz web es el punto de acceso principal de los usuarios al sistema. Debe ser intuitiva, responsiva y segura, con capacidad de gestionar todos los módulos de la plataforma desde un único entorno visual.")
    add_heading2(doc, "4.1 Requerimientos Funcionales de la Interfaz Web")

    add_section_label(doc, "RF-IW-01 · Autenticación y control de acceso")
    for txt in ["Login seguro con usuario y contraseña, con soporte para autenticación de dos factores (2FA)",
                "Gestión de sesiones con expiración automática por inactividad",
                "Recuperación de contraseña vía correo electrónico con token temporal",
                "Registro de intentos de acceso fallidos y bloqueo automático tras N intentos",
                "Soporte para SSO (Single Sign-On) con proveedores corporativos si se requiere"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-IW-02 · Panel de control (Dashboard)")
    for txt in ["Vista centralizada del estado del sistema: usuarios activos, formularios recibidos, alertas",
                "Gráficos y estadísticas en tiempo real sobre la actividad de la plataforma",
                "Acceso rápido a los módulos más utilizados según el rol del usuario",
                "Notificaciones internas del sistema (nuevos registros, errores, actualizaciones)"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-IW-03 · Gestión de usuarios desde la interfaz")
    for txt in ["Alta, baja y modificación de usuarios",
                "Asignación y modificación de roles y permisos",
                "Visualización del historial de actividad por usuario",
                "Exportación de listados de usuarios en formato Excel/PDF",
                "Envío de credenciales de acceso por correo electrónico desde la plataforma"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-IW-04 · Módulo de gestión documental")
    for txt in ["Carga, descarga y visualización de archivos (PDF, Excel, imágenes)",
                "Versionado de documentos con historial de cambios",
                "Clasificación por categorías y etiquetas",
                "Control de permisos de visualización y descarga por rol"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-IW-05 · Módulo de visualización de formularios recibidos")
    for txt in ["Listado de formularios recibidos desde usuarios remotos, ordenados por fecha/estado",
                "Filtros avanzados por fecha, tipo, estado y usuario",
                "Vista detallada de cada respuesta de formulario",
                "Validación manual y aprobación/rechazo de formularios desde la interfaz",
                "Exportación de datos de formularios a Excel o PDF"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-IW-06 · Módulo de reportes y auditoría")
    for txt in ["Generación de reportes parametrizables por rango de fechas, módulo y usuario",
                "Log de auditoría visible para administradores: quién hizo qué y cuándo",
                "Exportación de reportes en formato PDF y Excel",
                "Alertas configurables por correo ante eventos críticos del sistema"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-IW-07 · Diseño y accesibilidad")
    for txt in ["Interfaz responsiva: compatible con navegadores de escritorio y dispositivos móviles",
                "Compatibilidad con Chrome 100+, Firefox 100+, Edge 100+, Safari 15+",
                "Carga de páginas optimizada (< 3 segundos en conexión estándar)",
                "Manejo de errores amigable: mensajes claros, sin exposición de datos técnicos internos"]:
        add_bullet(doc, txt)
    add_divider(doc)

    # 5. FORMULARIOS
    insert_page_break(doc)
    add_heading1(doc, "5. Requerimientos de Formularios Independientes y Remotos")
    add_paragraph(doc, "Los formularios representan uno de los componentes más críticos del sistema, dado que deben operar de manera completamente independiente de la interfaz web principal. Su diseño garantiza que usuarios remotos puedan capturar y enviar información al sistema sin necesidad de tener acceso a la plataforma central.")
    add_heading2(doc, "5.1 Principios de Diseño de los Formularios")
    t4 = doc.add_table(rows=0, cols=2)
    set_table_width(t4, 9026)
    make_header_row(t4, ["Principio", "Descripción"], [2800, 6226])
    for i, row in enumerate([
        ("Independencia", "El formulario funciona como un módulo autónomo, alojado en URL propia, sin requerir login a la interfaz principal"),
        ("Operación remota", "Accesible desde cualquier dispositivo con conexión a internet, mediante enlace o código QR"),
        ("Persistencia de datos", "Cada envío es registrado directamente en la base de datos central a través de la API"),
        ("Modo offline / cola", "En caso de fallo de conectividad, los datos quedan en cola local y se sincronizan al recuperar la conexión"),
        ("Seguridad de envío", "Protección con token de sesión de formulario (no requiere cuenta de usuario)"),
        ("Validación en cliente", "Validación de campos obligatorios, formatos y rangos antes del envío"),
    ]):
        make_data_row(t4, list(row), [2800, 6226], shade=WHITE if i%2==0 else LIGHT_GRAY)

    add_heading2(doc, "5.2 Requerimientos Funcionales de Formularios")

    add_section_label(doc, "RF-FM-01 · Acceso y distribución")
    for txt in ["Formularios accesibles mediante URL única por tipo de formulario",
                "Distribución por código QR, enlace directo o correo electrónico",
                "Sin requerimiento de cuenta registrada para el usuario final del formulario",
                "Token de autenticación ligero para validar origen del envío"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-FM-02 · Tipos de formularios requeridos")
    t5 = doc.add_table(rows=0, cols=3)
    set_table_width(t5, 9026)
    make_header_row(t5, ["Tipo de Formulario", "Propósito", "Destino en BD"], [2500, 3000, 3526])
    for i, row in enumerate([
        ("Registro de novedades", "Captura de eventos operativos o incidencias", "tabla: respuestas_formulario"),
        ("Solicitud de información", "Peticiones de consulta o requerimientos", "tabla: solicitudes"),
        ("Actualización de datos", "Modificación de datos propios del usuario", "tabla: usuarios (pendiente aprobación)"),
        ("Reporte de actividad", "Informe periódico de actividades realizadas", "tabla: reportes_actividad"),
        ("Formulario de configuración", "Parametrización de opciones del usuario", "tabla: configuracion_usuario"),
    ]):
        make_data_row(t5, list(row), [2500, 3000, 3526], shade=WHITE if i%2==0 else LIGHT_GRAY)

    add_section_label(doc, "RF-FM-03 · Generación automática de documentos desde formularios")
    for txt in ["Al recibir un formulario, el sistema debe generar automáticamente un documento de confirmación (PDF)",
                "El documento generado debe quedar almacenado en la BD y ser descargable desde la interfaz web",
                "Notificación automática por correo al administrador cuando se recibe un formulario nuevo",
                "El usuario que envía recibe un acuse de recibo digital (correo o pantalla de confirmación)"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-FM-04 · Funcionamiento sin dependencia de la interfaz")
    for txt in ["El módulo de formularios tiene su propio proceso de servidor (microservicio o worker independiente)",
                "En caso de que la interfaz web principal esté caída, los formularios continúan recibiendo envíos",
                "Cola de procesamiento asíncrono: los datos llegan a BD independientemente del estado de la interfaz",
                "Panel de administración mínimo embebido para consulta básica sin acceder a la interfaz principal"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-FM-05 · Seguridad y validación de formularios")
    for txt in ["Protección contra spam mediante CAPTCHA o limitación de envíos por IP",
                "Validación de tipos de dato, longitud máxima y campos obligatorios",
                "Los formularios deben servirse sobre HTTPS obligatoriamente",
                "No almacenar datos sensibles en el navegador del usuario"]:
        add_bullet(doc, txt)
    add_divider(doc)

    # 6. GESTOR DE INTERFAZ
    insert_page_break(doc)
    add_heading1(doc, "6. Requerimientos del Gestor de Interfaz")
    add_paragraph(doc, "El gestor de interfaz es el módulo de administración técnica de la plataforma, destinado a los perfiles con rol de administrador o superadministrador. Permite configurar el comportamiento del sistema, gestionar la infraestructura técnica y supervisar el estado operativo de todos los componentes.")

    add_section_label(doc, "RF-GI-01 · Panel de administración del sistema")
    for txt in ["Vista de estado de todos los servicios activos (base de datos, API, formularios, interfaz)",
                "Indicadores de salud del sistema: CPU, memoria, almacenamiento y latencia de BD",
                "Gestión de versiones desplegadas y opción de rollback controlado",
                "Control de actualizaciones del sistema desde el panel"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-GI-02 · Gestión de configuración")
    for txt in ["Parámetros de configuración del sistema editables desde la interfaz (sin modificar archivos de servidor)",
                "Gestión de variables de entorno a través de panel seguro",
                "Configuración de plantillas de correo electrónico automático",
                "Definición de reglas de validación para formularios desde el gestor"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-GI-03 · Gestión de roles y permisos avanzada")
    for txt in ["Creación, edición y eliminación de roles personalizados",
                "Asignación granular de permisos por módulo y acción (ver, crear, editar, eliminar)",
                "Matriz de permisos visual para revisión rápida",
                "Historial de cambios en roles y permisos con autoría"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-GI-04 · Monitoreo y alertas")
    for txt in ["Registro de logs del sistema accesible desde el gestor",
                "Alertas automáticas por correo o notificación interna ante errores críticos, accesos no autorizados o caída de servicios",
                "Dashboard de métricas: formularios recibidos por día, usuarios activos, tiempo de respuesta API",
                "Reporte diario automático de estado del sistema enviado al administrador"]:
        add_bullet(doc, txt)

    add_section_label(doc, "RF-GI-05 · Gestión de copias de respaldo")
    for txt in ["Programación de backups automáticos desde el gestor",
                "Descarga manual de backups en formato comprimido (.sql.gz o similar)",
                "Restauración de backups con confirmación de doble paso",
                "Log de todas las operaciones de backup/restore realizadas"]:
        add_bullet(doc, txt)
    add_divider(doc)

    # 7. INFRAESTRUCTURA
    insert_page_break(doc)
    add_heading1(doc, "7. Requerimientos de Infraestructura y Despliegue")
    t6 = doc.add_table(rows=0, cols=2)
    set_table_width(t6, 9026)
    make_header_row(t6, ["Componente", "Especificación Mínima Requerida"], [2800, 6226])
    for i, row in enumerate([
        ("Servidor de aplicación", "VPS privado: 4 vCPU, 8 GB RAM, 100 GB SSD NVMe"),
        ("Sistema operativo", "Ubuntu 22.04 LTS o Debian 12 (soporte extendido)"),
        ("Base de datos", "PostgreSQL 15+ / MySQL 8+ en instancia dedicada o contenedor"),
        ("Servidor web / proxy inverso", "Nginx 1.24+ con configuración SSL/TLS Let's Encrypt o certificado privado"),
        ("Gestor de procesos", "PM2 (Node.js) o Gunicorn/uWSGI (Python) con monitoreo integrado"),
        ("Contenedores (opcional)", "Docker + Docker Compose para despliegue modular y reproducible"),
        ("Cola de mensajes", "RabbitMQ o Redis para procesamiento asíncrono de formularios"),
        ("Certificado SSL", "TLS 1.2/1.3 obligatorio; HTTPS forzado en todos los endpoints"),
        ("Correo transaccional", "SMTP propio o servicio (SendGrid/Mailgun) para notificaciones automáticas"),
        ("Control de versiones", "Git con repositorio privado (GitHub/GitLab)"),
        ("Dominio", "Dominio propio o subdominio configurado con DNS privado"),
        ("Disponibilidad objetivo", "99.5% mensual (aproximadamente < 4 horas de inactividad/mes)"),
    ]):
        make_data_row(t6, list(row), [2800, 6226], shade=WHITE if i%2==0 else LIGHT_GRAY)
    add_divider(doc)

    # 8. FLUJO GENERAL
    insert_page_break(doc)
    add_heading1(doc, "8. Flujo General de Operación del Sistema")
    add_paragraph(doc, "El flujo de operación de MineDax 2026 se organiza en tres caminos principales según el tipo de actor:")
    add_heading2(doc, "8.1 Flujo del Usuario de la Interfaz Web")
    for txt in ["El usuario accede a la URL de la plataforma desde su navegador",
                "Ingresa sus credenciales → el sistema valida contra la BD → emite token JWT",
                "Accede al dashboard según su rol (ADMIN, OPERATOR, VIEWER)",
                "Consulta formularios recibidos, genera reportes, gestiona usuarios o documentos",
                "Cierre de sesión o expiración automática → invalidación del token"]:
        add_bullet(doc, txt)
    add_heading2(doc, "8.2 Flujo del Usuario de Formulario Remoto")
    for txt in ["El usuario recibe enlace o QR del formulario correspondiente",
                "Accede al formulario desde cualquier dispositivo sin necesidad de login",
                "Completa los campos → el formulario valida en cliente → envía datos a la API",
                "La API valida, persiste en base de datos y emite confirmación",
                "El usuario recibe pantalla/correo de confirmación con número de radicado",
                "El administrador recibe notificación del nuevo envío en la interfaz web"]:
        add_bullet(doc, txt)
    add_heading2(doc, "8.3 Flujo del Administrador del Sistema (Gestor de Interfaz)")
    for txt in ["El SUPERADMIN accede al panel de administración con credenciales de alto privilegio",
                "Monitorea el estado de los servicios y revisa logs del sistema",
                "Gestiona usuarios, roles y configuraciones globales",
                "Programa y supervisa backups, actualizaciones y restauraciones",
                "Recibe y atiende alertas automáticas ante eventos críticos"]:
        add_bullet(doc, txt)
    add_divider(doc)

    # 9. RESUMEN
    insert_page_break(doc)
    add_heading1(doc, "9. Resumen de Requerimientos por Componente")
    t7 = doc.add_table(rows=0, cols=4)
    set_table_width(t7, 9026)
    make_header_row(t7, ["Código", "Componente", "Descripción Resumida", "Prioridad"], [1800, 2000, 3226, 2000])
    for i, row in enumerate([
        ("RF-DB-01", "Base de datos", "Disponibilidad 24/7, ACID, backups automáticos", "CRITICA"),
        ("RF-DB-02", "Base de datos", "Esquema de entidades definido y normalizado", "CRITICA"),
        ("RF-DB-03", "Base de datos", "Roles de usuario diferenciados con permisos específicos", "ALTA"),
        ("RF-DB-04", "Base de datos", "API de sincronización para formularios remotos", "CRITICA"),
        ("RF-DB-05", "Base de datos", "Cifrado TLS, hash de contraseñas, acceso restringido", "CRITICA"),
        ("RF-IW-01", "Interfaz Web", "Login seguro con 2FA y gestión de sesiones", "CRITICA"),
        ("RF-IW-02", "Interfaz Web", "Dashboard en tiempo real con estadísticas", "ALTA"),
        ("RF-IW-03", "Interfaz Web", "Gestión completa de usuarios desde la UI", "ALTA"),
        ("RF-IW-04", "Interfaz Web", "Módulo documental con versionado y permisos", "MEDIA"),
        ("RF-IW-05", "Interfaz Web", "Visualización y gestión de formularios recibidos", "ALTA"),
        ("RF-IW-06", "Interfaz Web", "Reportes parametrizables y log de auditoría", "ALTA"),
        ("RF-FM-01", "Formularios", "Acceso por URL/QR sin cuenta de usuario", "CRITICA"),
        ("RF-FM-02", "Formularios", "Tipos de formulario definidos con destino en BD", "CRITICA"),
        ("RF-FM-03", "Formularios", "Generación automática de documentos y notificaciones", "ALTA"),
        ("RF-FM-04", "Formularios", "Operación independiente de la interfaz web", "CRITICA"),
        ("RF-FM-05", "Formularios", "CAPTCHA, HTTPS, validación y seguridad", "ALTA"),
        ("RF-GI-01", "Gestor Interfaz", "Panel de estado de todos los servicios", "ALTA"),
        ("RF-GI-02", "Gestor Interfaz", "Configuración del sistema desde panel seguro", "ALTA"),
        ("RF-GI-03", "Gestor Interfaz", "Roles y permisos granulares con historial", "ALTA"),
        ("RF-GI-04", "Gestor Interfaz", "Monitoreo, logs y alertas automáticas", "ALTA"),
        ("RF-GI-05", "Gestor Interfaz", "Backups programados, descarga y restauración", "CRITICA"),
    ]):
        make_data_row(t7, list(row), [1800, 2000, 3226, 2000], shade=WHITE if i%2==0 else LIGHT_GRAY)

    doc.add_paragraph()
    add_paragraph(doc, "Este documento debe ser revisado, complementado con capturas de pantalla del sistema actual, y validado por el equipo de desarrollo antes de iniciar la fase de implementación en producción.", italic=True, color="555555")

    out = os.path.join(OUTPUT_DIR, "InformeTecnico_Requerimientos_MineDax2026.docx")
    doc.save(out)
    print(f"OK: InformeTecnico_Requerimientos_MineDax2026.docx  ({os.path.getsize(out):,} bytes)")
    return out


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 2: MENSAJE INFORMATIVO
# ═══════════════════════════════════════════════════════════════════════════════
def build_mensaje_informativo():
    doc = Document()
    set_page_margins(doc)

    # ENCABEZADO
    pb = doc.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pb.paragraph_format.space_before = Pt(10)
    pb.paragraph_format.space_after = Pt(0)
    pPr = pb._p.get_or_add_pPr()
    s = OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'), ACCENT_BLUE)
    pPr.append(s)
    add_run(pb, "MineDax 2026", bold=True, color=WHITE, size_pt=26)

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(0)
    ps.paragraph_format.space_after = Pt(9)
    pPr2 = ps._p.get_or_add_pPr()
    s2 = OxmlElement('w:shd'); s2.set(qn('w:val'),'clear'); s2.set(qn('w:color'),'auto'); s2.set(qn('w:fill'), ACCENT_BLUE)
    pPr2.append(s2)
    add_run(ps, "Plataforma de Gestión en Línea — Línea Privada", color="BDD7EE", size_pt=13)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_after = Pt(3)
    add_run(pt, "MENSAJE INFORMATIVO DE REQUERIMIENTOS PARA EL EQUIPO DE DESARROLLO", bold=True, color=ACCENT_BLUE, size_pt=13)

    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd.paragraph_format.space_after = Pt(3)
    add_run(pd, "Mayo 2026  ·  Versión Resumida  ·  Sujeto a complemento con imágenes", italic=True, color="777777", size_pt=9.5)

    add_divider(doc)

    # CONTEXTO
    add_paragraph(doc, "El presente mensaje tiene por objetivo informar al equipo de desarrollo, de manera clara y directa, sobre los requerimientos esenciales que debe cumplir la plataforma MineDax 2026 para operar correctamente en línea a través de una línea privada. La plataforma integra tres componentes interdependientes: base de datos en línea, interfaz web y formularios remotos independientes.")
    add_paragraph(doc, "Este documento es complementario al Informe Técnico de Requerimientos Funcionales y está pensado como guía de referencia rápida para priorizar la implementación.")
    add_divider(doc)

    # 1. POR QUÉ SE NECESITA
    add_heading1(doc, "1. ¿Por Qué Se Necesita Este Sistema?")
    add_callout(doc, "Problema central que resuelve MineDax 2026",
        "La organización carece de una plataforma unificada para gestionar, almacenar y consultar información operativa de forma continua y remota. Los procesos actuales dependen de herramientas desconectadas, lo que genera pérdida de información, duplicación de esfuerzos y ausencia de trazabilidad.",
        MID_BLUE, LIGHT_BLUE, ACCENT_BLUE)
    doc.add_paragraph()
    t_why = doc.add_table(rows=0, cols=2)
    set_table_width(t_why, 9026)
    make_header_row(t_why, ["Necesidad", "Argumento"], [3200, 5826])
    for i, row in enumerate([
        ("Centralización de datos", "Sin una BD centralizada en línea, cada área trabaja en silos, imposibilitando reportes unificados y auditoría real del sistema"),
        ("Continuidad operativa", "El sistema debe estar disponible 24/7 para garantizar que las operaciones no se detengan por indisponibilidad tecnológica"),
        ("Captura remota de información", "Los formularios independientes permiten registrar datos desde campo o ubicaciones remotas, sin depender de acceso a la plataforma principal"),
        ("Control de acceso y trazabilidad", "La gestión de roles y el log de auditoría son indispensables para controlar quién accede y qué modifica"),
        ("Gestión administrativa eficiente", "El gestor de interfaz elimina la necesidad de acceder al servidor directamente para tareas de configuración, reduciendo riesgos operativos"),
    ]):
        make_data_row(t_why, list(row), [3200, 5826], shade=WHITE if i%2==0 else LT_GRAY2)
    add_divider(doc)

    # 2. PUNTOS CLAVE
    add_heading1(doc, "2. Puntos Clave por Componente")

    add_heading2(doc, "BASE DE DATOS")
    add_callout(doc, "¿Qué se requiere?",
        "Una base de datos relacional (PostgreSQL o MySQL) activa de forma continua, con backups automáticos, replicación, cifrado en reposo y conexiones SSL. El acceso debe estar restringido a la capa de API, nunca expuesto directamente a internet.",
        MID_BLUE, LIGHT_BLUE, ACCENT_BLUE)
    for txt in ["Motor: PostgreSQL 15+ o MySQL 8+",
                "Disponibilidad: 24/7 — operación continua sin interrupciones programadas",
                "Backups: diarios completos + incrementales por hora",
                "Cifrado: TLS en tránsito, AES-256 en reposo para datos sensibles",
                "Roles de acceso: SUPERADMIN, ADMIN, OPERATOR, VIEWER, FORM_USER"]:
        add_bullet(doc, txt)
    add_callout(doc, "¿Por qué es crítico?",
        "Sin una base de datos robusta y siempre activa, toda la información del sistema se vuelve volátil. La pérdida de datos operativos tiene impacto directo en la continuidad del negocio y en la confianza del equipo en la plataforma.",
        LT_ORANGE, ORANGE, ORANGE)
    doc.add_paragraph()

    add_heading2(doc, "INTERFAZ WEB")
    add_callout(doc, "¿Qué se requiere?",
        "Una aplicación web responsiva (React.js / Next.js recomendado) con login seguro, panel de control, módulos de gestión de usuarios, documentos, formularios recibidos y reportes. Compatible con los principales navegadores y optimizada para tiempos de carga bajos.",
        MID_BLUE, LIGHT_BLUE, ACCENT_BLUE)
    for txt in ["Autenticación con soporte 2FA y gestión de sesiones con expiración automática",
                "Dashboard en tiempo real: usuarios activos, formularios nuevos, alertas del sistema",
                "Módulo de usuarios: alta, baja, edición, asignación de roles",
                "Módulo de formularios: visualización, validación y exportación de respuestas",
                "Módulo de reportes: parametrizables, exportables en PDF y Excel",
                "Módulo documental: carga, versionado, descarga controlada por rol"]:
        add_bullet(doc, txt)
    add_callout(doc, "¿Por qué es crítico?",
        "La interfaz es el punto de interacción principal del equipo gestor con el sistema. Sin ella debidamente configurada, los administradores no pueden supervisar, validar ni actuar sobre la información recibida desde los formularios remotos.",
        LT_ORANGE, ORANGE, ORANGE)
    doc.add_paragraph()

    add_heading2(doc, "FORMULARIOS INDEPENDIENTES Y REMOTOS")
    add_callout(doc, "¿Qué se requiere?",
        "Formularios web autónomos accesibles mediante URL o código QR, que no requieran login de usuario, funcionen desde cualquier dispositivo con internet, envíen datos directamente a la base de datos a través de la API, y operen incluso si la interfaz web principal no está disponible.",
        MID_BLUE, LIGHT_BLUE, ACCENT_BLUE)
    for txt in ["Acceso por URL única o QR — sin cuenta de usuario requerida",
                "Envío directo a BD vía API con validación de datos en cliente y servidor",
                "Modo de cola: si hay fallo de conectividad, los datos se guardan y sincronizan al recuperar conexión",
                "Generación automática de documento de confirmación (PDF) por cada envío",
                "Notificación automática al administrador ante nuevo formulario recibido",
                "Protección: CAPTCHA + HTTPS obligatorio + validación de campos"]:
        add_bullet(doc, txt)
    add_callout(doc, "¿Por qué es crítico?",
        "Los formularios son el canal principal de entrada de información al sistema desde usuarios en campo. Si dependen de la interfaz, cualquier caída del servidor dejaría sin capacidad de registro a todos los usuarios remotos, generando pérdida de datos en tiempo real.",
        LT_ORANGE, ORANGE, ORANGE)

    insert_page_break(doc)

    add_heading2(doc, "GESTOR DE INTERFAZ (ADMINISTRACIÓN TÉCNICA)")
    add_callout(doc, "¿Qué se requiere?",
        "Un panel de administración técnica para el SUPERADMIN que permita monitorear el estado de todos los servicios, gestionar configuraciones del sistema, administrar roles y permisos, programar backups, revisar logs y recibir alertas automáticas ante eventos críticos.",
        MID_BLUE, LIGHT_BLUE, ACCENT_BLUE)
    for txt in ["Panel de estado: salud de servicios (BD, API, formularios, interfaz)",
                "Configuración del sistema desde interfaz segura sin acceso directo al servidor",
                "Gestión de roles con matriz de permisos granular y auditable",
                "Alertas automáticas: correo/notificación ante errores críticos o accesos no autorizados",
                "Gestión de backups: programación, descarga y restauración con confirmación"]:
        add_bullet(doc, txt)
    add_callout(doc, "¿Por qué es crítico?",
        "Sin un gestor de interfaz, las tareas de mantenimiento del sistema requieren acceso directo al servidor, lo que aumenta exponencialmente el riesgo de errores humanos y compromete la seguridad de la plataforma.",
        LT_ORANGE, ORANGE, ORANGE)
    add_divider(doc)

    # 3. FLUJO DE TRABAJO
    add_heading1(doc, "3. Flujo de Trabajo del Sistema")
    add_paragraph(doc, "El siguiente esquema resume cómo fluye la información a través de los componentes del sistema:")
    t_flow = doc.add_table(rows=0, cols=4)
    set_table_width(t_flow, 9026)
    make_header_row(t_flow, ["#", "Actor", "Componente", "Acción Principal"], [400, 2200, 1800, 4626])
    for i, row in enumerate([
        ("1", "Usuario remoto", "Formulario", "Accede por URL/QR → llena datos → envía → recibe confirmación"),
        ("2", "API", "Base de Datos", "Recibe datos del formulario → valida → persiste en BD"),
        ("3", "BD", "Base de Datos", "Almacena registro → activa trigger de notificación"),
        ("4", "Administrador", "Interfaz Web", "Recibe alerta → consulta formulario recibido → valida/aprueba"),
        ("5", "Administrador", "Interfaz Web", "Genera reporte → exporta PDF/Excel → gestiona usuarios"),
        ("6", "SUPERADMIN", "Gestor Interfaz", "Monitorea servicios → configura sistema → gestiona backups"),
        ("7", "Sistema", "Gestor Interfaz", "Envía alertas automáticas → ejecuta backups programados"),
    ]):
        make_data_row(t_flow, list(row), [400, 2200, 1800, 4626], shade=WHITE if i%2==0 else LT_GRAY2)
    doc.add_paragraph()
    add_callout(doc, "Independencia de componentes — punto clave para el equipo de desarrollo",
        "Los formularios deben funcionar como microservicio independiente. Si la interfaz web principal cae, los formularios siguen recibiendo y persistiendo datos. La base de datos es la única dependencia compartida, y debe tener alta disponibilidad garantizada.",
        LT_GREEN, GREEN, GREEN)
    add_divider(doc)

    # 4. RESUMEN EJECUTIVO
    add_heading1(doc, "4. Resumen Ejecutivo — Lo Esencial Sin Excepción")
    t_sum = doc.add_table(rows=0, cols=4)
    set_table_width(t_sum, 9026)
    make_header_row(t_sum, ["#", "Requerimiento Clave", "Componente", "Impacto si no se implementa"], [400, 3000, 1600, 4026])
    for i, row in enumerate([
        ("1",  "BD activa 24/7 con backups automáticos",          "Base de datos",  "Pérdida de datos, inoperabilidad del sistema"),
        ("2",  "API segura para recepción de formularios remotos", "Base de datos",  "Formularios no pueden guardar datos sin depender de la UI"),
        ("3",  "Cifrado TLS en tránsito y AES en reposo",         "Base de datos",  "Exposición de datos sensibles y riesgo legal"),
        ("4",  "Login seguro con 2FA y gestión de sesiones",       "Interfaz Web",   "Acceso no autorizado a toda la plataforma"),
        ("5",  "Dashboard en tiempo real",                         "Interfaz Web",   "Sin visibilidad del estado del sistema ni de nuevos ingresos"),
        ("6",  "Formularios accesibles sin login por URL/QR",      "Formularios",    "Usuarios remotos no pueden ingresar información al sistema"),
        ("7",  "Formularios independientes de la interfaz web",    "Formularios",    "Caída de la UI deja sin captura de datos en campo"),
        ("8",  "Generación automática de confirmación por envío",  "Formularios",    "Sin trazabilidad ni respaldo del dato ingresado"),
        ("9",  "Panel de estado de todos los servicios",           "Gestor Interfaz","El SUPERADMIN no puede actuar ante fallas sin acceso al servidor"),
        ("10", "Gestión de roles y permisos granular",             "Gestor Interfaz","Sin control de acceso diferenciado, todos los usuarios ven todo"),
    ]):
        make_data_row(t_sum, list(row), [400, 3000, 1600, 4026], shade=WHITE if i%2==0 else LT_GRAY2)

    doc.add_paragraph()
    add_callout(doc, "Nota para el equipo de desarrollo",
        "Este documento será complementado próximamente con capturas de pantalla del sistema en su estado actual. Se solicita al equipo indicar en qué servidor o plataforma cloud se puede alojar el sistema, de acuerdo con las especificaciones de infraestructura descritas en el informe técnico adjunto.",
        MID_BLUE, LIGHT_BLUE, ACCENT_BLUE)
    doc.add_paragraph()
    add_paragraph(doc, "Cualquier duda o solicitud de aclaración sobre los requerimientos puede dirigirse al responsable del proyecto a través de la mesa de ayuda. Se agradece la revisión oportuna de ambos documentos para agendar la reunión de planificación técnica.", italic=True, color="555555")

    out = os.path.join(OUTPUT_DIR, "MensajeInformativo_Requerimientos_MineDax2026.docx")
    doc.save(out)
    print(f"OK: MensajeInformativo_Requerimientos_MineDax2026.docx  ({os.path.getsize(out):,} bytes)")
    return out


# ─── MAIN ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("  Generando documentos MineDax 2026")
    print("=" * 60)
    try:
        build_informe_tecnico()
    except Exception as e:
        print(f"ERROR en Informe Tecnico: {e}")
        import traceback; traceback.print_exc()
    try:
        build_mensaje_informativo()
    except Exception as e:
        print(f"ERROR en Mensaje Informativo: {e}")
        import traceback; traceback.print_exc()
    print("=" * 60)
    print("  Listo. Archivos guardados en la misma carpeta.")
    print("=" * 60)
